from docx import Document

# Create a new Document
doc = Document()

# Title
doc.add_heading('6. Merkle-Based Post-Quantum OT: Mathematical Framework', level=1)

# Section 6.1
doc.add_heading('6.1 Key Definitions', level=2)

# Security Parameters
doc.add_paragraph('Security Parameters:')
doc.add_paragraph('    n: Number of secrets held by the Sender.')
doc.add_paragraph('    k: Security parameter (e.g., the bit-length of hashes and salts).')

# Merkle Tree
doc.add_paragraph('Merkle Tree:')
doc.add_paragraph('    Leaves: L = { Hk(m1), Hk(m2), ..., Hk(mn) }, where Hk(x) is a quantum-safe hash function (e.g., SHA-3) and mi are the secrets.')
doc.add_paragraph('    Tree Construction: Internal nodes are defined as Hk(Li || Lj), where Li, Lj are sibling nodes, and || denotes concatenation.')
doc.add_paragraph('    Root: The root R is the top-most hash of the tree, uniquely identifying the set of secrets.')

# Merkle Proof
doc.add_paragraph('Merkle Proof:')
doc.add_paragraph('    A proof Pi for leaf i is a sequence of hashes that reconstructs R when combined with Hk(mi).')

# Section 6.2
doc.add_heading('6.2 Protocol Phases', level=2)

# Phase 1: Setup by Sender
doc.add_heading('Phase 1: Setup by Sender', level=3)

# Input
doc.add_paragraph('Input:')
doc.add_paragraph('    M = { m1, m2, ..., mn }: The secrets.')
doc.add_paragraph('    Hk: A quantum-safe hash function.')

# Output
doc.add_paragraph('Output:')
doc.add_paragraph('    R: The root of the Merkle tree.')
doc.add_paragraph('    Pi: The Merkle proof for each i in {1, 2, ..., n}.')

# Procedure
doc.add_paragraph('Procedure:')
doc.add_paragraph('    1. Compute Hk(mi) for each secret mi.')
doc.add_paragraph('    2. Construct a Merkle tree from { Hk(m1), ..., Hk(mn) }.')
doc.add_paragraph('    3. Share R with the Receiver.')

# Phase 2: Oblivious Transfer
doc.add_heading('Phase 2: Oblivious Transfer', level=3)

# Receiver's Selection
doc.add_paragraph('Receiver’s Selection:')
doc.add_paragraph('    i: The index of the desired secret.')
doc.add_paragraph('    Generate a query mask Q such that Q(i) is indistinguishable from random noise.')

# Sender's Response
doc.add_paragraph('Sender’s Response:')
doc.add_paragraph('    Compute Pi: The Merkle proof for the i-th leaf.')
doc.add_paragraph('    Return { Hk(mi), Pi }.')

# Verification by Receiver
doc.add_paragraph('Verification by Receiver:')
doc.add_paragraph('    Verify the Merkle proof: R = Reconstruct(Hk(mi), Pi).')
doc.add_paragraph('    If valid, extract mi using additional context (e.g., decoding the hash).')

# Section 6.3
doc.add_heading('6.3 Properties', level=2)

# Correctness
doc.add_paragraph('Correctness:')
doc.add_paragraph('    If the Receiver follows the protocol, they will retrieve mi such that:')
doc.add_paragraph('        Hk(mi) is valid under Pi → R.')

# Privacy
doc.add_paragraph('Privacy:')
doc.add_paragraph('    The Sender cannot determine i due to the masking of the query Q(i).')

# Quantum Resistance
doc.add_paragraph('Quantum Resistance:')
doc.add_paragraph('    The security relies on the hardness of inverting Hk, a quantum-safe hash function.')

# Section 6.4
doc.add_heading('6.4 Optional Extensions', level=2)

# Salting for Hk(mi)
doc.add_paragraph('Add Salting for Hk(mi):')
doc.add_paragraph('    Hk(mi || si), where si is a unique salt per secret to prevent replay attacks.')

# Adaptive Masking for Q(i)
doc.add_paragraph('Use Adaptive Masking for Q(i):')
doc.add_paragraph('    Ensuring it resists statistical attacks.')

# Save the document
doc.save('Merkle_Based_Post_Quantum_OT_Mathematical_Framework.docx')

print("Document created successfully!")
